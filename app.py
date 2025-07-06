import os
import io
import fitz  # PyMuPDF
from flask import Flask, request, render_template, send_file, flash, redirect
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

def convert_pdf_to_docx(pdf_path, output_stream):
    """
    Converts a PDF to DOCX, with a protection mechanism against creating
    blank pages due to large vertical gaps in the PDF.
    """
    try:
        pdf_document = fitz.open(pdf_path)
        word_document = Document()
        # Set document margins
        for section in word_document.sections:
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)

        for page_num in range(len(pdf_document)):
            
            # Correct page break logic
            if page_num > 0:
                word_document.add_page_break()

            page = pdf_document.load_page(page_num)
            
            blocks = page.get_text("dict", flags=fitz.TEXTFLAGS_DICT, sort=True)["blocks"]
            if not blocks: continue

            try:
                base_x0 = min(block['bbox'][0] for block in blocks if block['lines'])
            except ValueError:
                base_x0 = 0

            last_block_y1 = 0.0
            last_font_size = 12.0

            for block in blocks:
                block_text = "".join(span["text"] for line in block.get("lines", []) for span in line.get("spans", [])).strip()
                if not block_text:
                    continue # Skip empty blocks

                p = word_document.add_paragraph()
                p_format = p.paragraph_format

                # --- Protection mechanism against large spaces ---
                current_y0 = block['bbox'][1]
                vertical_gap = current_y0 - last_block_y1
                normal_line_gap = last_font_size * 1.2
                
                extra_space = vertical_gap - normal_line_gap
                # Only add significant, but capped, extra spacing
                if 5 < extra_space < 100: # Cap prevents huge gaps from creating blank pages
                    p_format.space_before = Pt(extra_space)
                # --- End of protection mechanism ---

                current_x0 = block['bbox'][0]
                indentation = current_x0 - base_x0
                if indentation > 1:
                    p_format.left_indent = Pt(indentation)
                
                p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

                for line in block["lines"]:
                    for span in line["spans"]:
                        run = p.add_run(span["text"])
                        font_name = span['font'].split('-')[0]
                        run.font.name = font_name
                        run.font.size = Pt(int(span['size'] + 0.5))
                        last_font_size = run.font.size.pt
                        
                        if "bold" in span['font'].lower() or (span['flags'] & 16):
                            run.bold = True
                        if "italic" in span['font'].lower() or (span['flags'] & 2):
                            run.italic = True
                
                last_block_y1 = block['bbox'][3]

        word_document.save(output_stream)
        return True
    except Exception as e:
        print(f"Error during conversion: {e}")
        return False


# --- Flask routes remain unchanged ---
app = Flask(__name__)
app.config['SECRET_KEY'] = 'a-very-secret-key'
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_and_convert():
    if 'pdf_file' not in request.files: flash('No file part'); return redirect(request.url)
    file = request.files['pdf_file']
    if file.filename == '': flash('No selected file'); return redirect(request.url)
    if file and file.filename.endswith('.pdf'):
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        try:
            file.save(pdf_path)
            file_stream = io.BytesIO()
            success = convert_pdf_to_docx(pdf_path, file_stream)
            file_stream.seek(0)
            if success:
                docx_filename = os.path.splitext(file.filename)[0] + '.docx'
                return send_file(file_stream, as_attachment=True, download_name=docx_filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            else:
                flash('An error occurred during conversion.'); return redirect(request.url)
        finally:
            if os.path.exists(pdf_path): os.remove(pdf_path)
    else:
        flash('Invalid file type. Please upload a PDF.'); return redirect(request.url)

if __name__ == '__main__':
    app.run(debug=True)